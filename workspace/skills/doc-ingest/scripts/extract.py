#!/usr/bin/env python3
"""
Document extraction script.
Extracts text, tables, and metadata from PDF, Excel, Word, and CSV files.

Usage:
    python3 extract.py <input_file> <output_json>
"""

import sys
import json
import os
from datetime import datetime, timezone
from pathlib import Path


def extract_pdf(filepath):
    """Extract text and tables from PDF using pdfplumber."""
    import pdfplumber

    chunks = []
    tables = []

    with pdfplumber.open(filepath) as pdf:
        num_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                chunks.append({
                    "page": i + 1,
                    "text": text.strip(),
                    "type": "text"
                })

            page_tables = page.extract_tables()
            for j, table in enumerate(page_tables):
                if table and len(table) > 1:
                    headers = table[0] if table[0] else [f"col_{k}" for k in range(len(table[0]))]
                    rows = []
                    for row in table[1:]:
                        row_dict = {}
                        for k, cell in enumerate(row):
                            key = headers[k] if k < len(headers) and headers[k] else f"col_{k}"
                            row_dict[key] = cell
                        rows.append(row_dict)

                    tables.append({
                        "page": i + 1,
                        "table_index": j,
                        "headers": headers,
                        "rows": rows
                    })

    return {
        "format": "pdf",
        "pages": num_pages,
        "chunks": chunks,
        "tables": tables
    }


def extract_excel(filepath):
    """Extract data from Excel files using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(filepath, read_only=True, data_only=True)
    chunks = []
    tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = []
        headers = None

        for i, row in enumerate(ws.iter_rows(values_only=True)):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if i == 0:
                headers = row_values
                continue
            if any(v.strip() for v in row_values):
                row_dict = {}
                for j, val in enumerate(row_values):
                    key = headers[j] if headers and j < len(headers) else f"col_{j}"
                    row_dict[key] = val
                rows_data.append(row_dict)

        if rows_data:
            tables.append({
                "sheet": sheet_name,
                "headers": headers,
                "rows": rows_data,
                "row_count": len(rows_data)
            })

            # Also create a text chunk for search
            text_lines = [f"Sheet: {sheet_name}"]
            if headers:
                text_lines.append("Columns: " + ", ".join(headers))
            text_lines.append(f"Rows: {len(rows_data)}")
            chunks.append({
                "sheet": sheet_name,
                "text": "\n".join(text_lines),
                "type": "spreadsheet"
            })

    wb.close()

    return {
        "format": "excel",
        "sheets": len(wb.sheetnames),
        "chunks": chunks,
        "tables": tables
    }


def extract_docx(filepath):
    """Extract text and tables from Word documents."""
    from docx import Document

    doc = Document(filepath)
    chunks = []
    tables = []

    # Extract paragraphs
    current_section = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_section:
                chunks.append({
                    "text": "\n".join(current_section),
                    "type": "text"
                })
                current_section = []
            continue
        current_section.append(text)

    if current_section:
        chunks.append({
            "text": "\n".join(current_section),
            "type": "text"
        })

    # Extract tables
    for i, table in enumerate(doc.tables):
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        rows_data = []
        for row in table.rows[1:]:
            row_dict = {}
            for j, cell in enumerate(row.cells):
                key = headers[j] if j < len(headers) else f"col_{j}"
                row_dict[key] = cell.text.strip()
            rows_data.append(row_dict)

        if rows_data:
            tables.append({
                "table_index": i,
                "headers": headers,
                "rows": rows_data
            })

    return {
        "format": "docx",
        "paragraphs": len(chunks),
        "chunks": chunks,
        "tables": tables
    }


def extract_csv(filepath):
    """Extract data from CSV files."""
    import csv

    chunks = []
    tables = []

    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        # Detect delimiter
        sample = f.read(4096)
        f.seek(0)
        dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')

        reader = csv.reader(f, dialect)
        headers = next(reader, None)

        rows_data = []
        for row in reader:
            if any(cell.strip() for cell in row):
                row_dict = {}
                for j, val in enumerate(row):
                    key = headers[j] if headers and j < len(headers) else f"col_{j}"
                    row_dict[key] = val
                rows_data.append(row_dict)

    if rows_data:
        tables.append({
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data)
        })

        text_lines = []
        if headers:
            text_lines.append("Columns: " + ", ".join(headers))
        text_lines.append(f"Rows: {len(rows_data)}")
        chunks.append({
            "text": "\n".join(text_lines),
            "type": "csv"
        })

    return {
        "format": "csv",
        "chunks": chunks,
        "tables": tables
    }


EXTRACTORS = {
    '.pdf': extract_pdf,
    '.xlsx': extract_excel,
    '.xls': extract_excel,
    '.docx': extract_docx,
    '.csv': extract_csv,
}


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 extract.py <input_file> <output_json>")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    ext = input_file.suffix.lower()
    if ext not in EXTRACTORS:
        print(f"Error: Unsupported format: {ext}")
        print(f"Supported: {', '.join(EXTRACTORS.keys())}")
        sys.exit(1)

    print(f"Extracting: {input_file.name} ({ext})")
    result = EXTRACTORS[ext](str(input_file))

    # Add metadata
    result["filename"] = input_file.name
    result["filesize_bytes"] = input_file.stat().st_size
    result["extracted_at"] = datetime.now(timezone.utc).isoformat()

    # Generate summary
    total_chunks = len(result.get("chunks", []))
    total_tables = len(result.get("tables", []))
    total_text = sum(len(c.get("text", "")) for c in result.get("chunks", []))
    result["summary"] = {
        "chunks": total_chunks,
        "tables": total_tables,
        "total_chars": total_text
    }

    # Write output
    output_file.parent.mkdir(parents=True, exist_ok=True)
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Done: {total_chunks} chunks, {total_tables} tables, {total_text} chars")
    print(f"Output: {output_file}")


if __name__ == "__main__":
    main()
